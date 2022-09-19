from docx import Document

document = Document()

document.add_heading("기사 제목", level = 0)
document.add_paragraph("기사 링크")
document.add_paragraph('기사 본문')

document.save("test.docx")
