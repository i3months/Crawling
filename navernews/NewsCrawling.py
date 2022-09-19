import requests 
from bs4 import BeautifulSoup
import time
import pyautogui
from docx import Document
from openpyxl import Workbook
from openpyxl.styles import Alignment

# user input
keyword = pyautogui.prompt("검색어를 입력하세요")
lastpage = int(pyautogui.prompt("몇 페이지까지 가져올까요?"))

# word document create
document = Document()

# create excel
wb = Workbook()
ws = wb.create_sheet(keyword)

# row width control
ws.column_dimensions['A'].width = 60
ws.column_dimensions['B'].width = 60
ws.column_dimensions['C'].width = 60


page_num = 1
row = 1

for i in range(1, lastpage * 10, 10):
    print(f"{page_num} 번째 페이지를 가져오고 있습니다.")
    response = requests.get(f"https://search.naver.com/search.naver?where=news&sm=tab_jum&query={keyword}&start={i}")
    html = response.text
    soup = BeautifulSoup(html, 'html.parser')

    articles = soup.select('div.info_group')

    for article in articles:
        links = article.select("a.info")
        if len(links) >= 2:
            url = links[1].attrs['href']
            response = requests.get(url, headers={'User-agent':'Mozila/5.0'})
            html = response.text
            soup_sub = BeautifulSoup(html, 'html.parser')                    

            if "entertain" in response.url:
                title = soup_sub.select_one(".end_tit")
                content = soup_sub.select_one("#articeBody")
            elif "sports" in response.url:
                title = soup_sub.select_one("h4.title")
                content = soup_sub.select_one("#newsEndContents")

                # delete unnecessary tag
                divs = content.select("div")
                for div in divs:
                    div.decompose()

                paragraphs = content.select("p")
                for paragraph in paragraphs:
                    paragraph.decompose()
            else:
                title = soup_sub.select_one("h2.media_end_head_headline")
                content = soup_sub.select_one("#dic_area")
            
            print("----link----\n", url)
            print("----title----\n", title.text.strip())
            print("----paragraph----\n", content.text.strip())

            document.add_heading(title.text.strip(), level = 0)
            document.add_paragraph(url)
            document.add_paragraph(content.text.strip())

            ws[f'A{row}'] = url
            ws[f'B{row}'] = title.text.strip()
            ws[f'C{row}'] = content.text.strip()

            # line control
            ws[f'C{row}'].alignment = Alignment(wrap_text=True)
            

            time.sleep(0.3)        
            row += 1    

    # last page check
    isLastPage = soup.select_one('a.btn_next').attr['aria-disabled']
    if isLastPage:
        print('마지막 페이지 입니다')
        break
    page_num += 1

document.save(f"{keyword}_result.docx")
wb.save(f'{keyword}_result.xlsx')
